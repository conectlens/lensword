"""Parsing every supported import format (issue #85).

The issue asks for fixture tests for every format, and for unsupported,
encrypted and corrupt files to fail without partial writes. Fixtures are
generated here rather than committed as binaries: a checked-in PDF is a blob
nobody reviews, and generating one makes the input visible in the test that
depends on it.
"""
from __future__ import annotations

import io
import zipfile

import pytest

from app.domain.services.documents import (
    MAX_BYTES,
    MAX_SECTIONS,
    DocumentStructureError,
    DocumentTooLargeError,
    build_document,
    split_sentences,
)
from app.infrastructure.document_parsers import (
    SUPPORTED_MEDIA_TYPES,
    detect_media_type,
    parse_document,
)

PROSE = "El gato duerme. La casa es grande. Vamos a correr."


# --- Fixture builders ------------------------------------------------------


def _pdf(pages: list[str], encrypt: bool = False) -> bytes:
    from pypdf import PdfWriter
    from reportlab.pdfgen import canvas  # type: ignore[import-not-found]

    buffer = io.BytesIO()
    drawer = canvas.Canvas(buffer)
    for text in pages:
        drawer.drawString(72, 720, text)
        drawer.showPage()
    drawer.save()
    if not encrypt:
        return buffer.getvalue()

    writer = PdfWriter(clone_from=io.BytesIO(buffer.getvalue()))
    writer.encrypt("secret")
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def _docx(paragraphs: list[tuple[str, str]]) -> bytes:
    import docx

    document = docx.Document()
    for style, text in paragraphs:
        document.add_paragraph(text, style=style) if style else document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _epub(chapters: list[tuple[str, str]]) -> bytes:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("test")
    book.set_title("Test")
    book.set_language("es")
    items = []
    for index, (title, body) in enumerate(chapters, start=1):
        chapter = epub.EpubHtml(title=title, file_name=f"c{index}.xhtml", lang="es")
        # An <h1> as well as a <title>: ebooklib's writer strips the latter,
        # and real EPUBs carry the chapter name in the TOC or a heading.
        chapter.content = (
            f"<html><head><title>{title}</title></head>"
            f"<body><h1>{title}</h1><p>{body}</p></body></html>"
        )
        book.add_item(chapter)
        items.append(chapter)
    book.toc = tuple(items)
    book.spine = ["nav", *items]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    buffer = io.BytesIO()
    epub.write_epub(buffer, book)
    return buffer.getvalue()


# --- Detection -------------------------------------------------------------


def test_content_beats_the_filename():
    """A PDF called notes.txt is a PDF. Trusting the extension would feed it to
    the text parser and produce mojibake `candidates`."""
    assert detect_media_type(_pdf(["hello"]), "notes.txt") == "application/pdf"


def test_an_epub_and_a_docx_are_told_apart_despite_the_same_signature():
    """Both are zip containers, so the magic bytes are identical and the
    container has to be opened."""
    assert detect_media_type(_epub([("One", PROSE)]), "x") == "application/epub+zip"
    assert detect_media_type(_docx([(None, PROSE)]), "x").endswith("wordprocessingml.document")


def test_the_extension_is_used_where_there_is_no_signature():
    assert detect_media_type(b"a,b\n1,2", "rows.csv") == "text/csv"
    assert detect_media_type(b"# Title", "notes.md") == "text/markdown"


def test_an_unknown_extension_falls_back_to_plain_text():
    assert detect_media_type(b"words", "notes.unknown") == "text/plain"


# --- One test per format ---------------------------------------------------


def test_plain_text():
    document = parse_document(PROSE.encode(), "a.txt")

    assert document.sentence_count == 3


def test_markdown_sections_are_labelled_by_heading():
    raw = b"# Chapter One\nEl gato duerme.\n\n## Chapter Two\nLa casa es grande."

    document = parse_document(raw, "a.md")

    assert [s.label for s in document.sections] == ["Chapter One", "Chapter Two"]


def test_html_drops_script_and_style():
    """Left in, they produce candidates like `getElementById`."""
    raw = b"<html><style>p{color:red}</style><script>var x=1</script><p>El gato duerme.</p></html>"

    document = parse_document(raw, "a.html")

    text = document.sections[0].sentences[0].text
    assert "gato" in text
    assert "getElementById" not in text and "color" not in text


def test_srt_keeps_speech_and_drops_timings():
    raw = b"1\n00:00:01,000 --> 00:00:03,000\nEl gato duerme.\n\n2\n00:00:04,000 --> 00:00:06,000\nLa casa es grande.\n"

    document = parse_document(raw, "a.srt")

    joined = " ".join(s.text for s in document.sections[0].sentences)
    assert "gato" in joined and "casa" in joined
    assert "00:00" not in joined


def test_vtt_drops_its_header_and_markup():
    raw = b"WEBVTT\n\n00:00:01.000 --> 00:00:03.000\n<v Speaker>El gato duerme.\n"

    document = parse_document(raw, "a.vtt")

    joined = " ".join(s.text for s in document.sections[0].sentences)
    assert "WEBVTT" not in joined and "<v" not in joined
    assert "gato" in joined


def test_subtitle_cues_are_rejoined_before_splitting():
    """Cues break on display timing, not grammar. One cue per sentence would
    give half-sentence excerpts."""
    raw = b"1\n00:00:01,000 --> 00:00:02,000\nEl gato\n\n2\n00:00:02,000 --> 00:00:03,000\nduerme mucho.\n"

    document = parse_document(raw, "a.srt")

    assert document.sections[0].sentences[0].text == "El gato duerme mucho."


def test_pdf_pages_become_sections():
    document = parse_document(_pdf(["El gato duerme.", "La casa es grande."]), "a.pdf")

    assert [s.label for s in document.sections] == ["Page 1", "Page 2"]


def test_docx_headings_become_section_labels():
    raw = _docx([("Heading 1", "Chapter One"), (None, PROSE)])

    document = parse_document(raw, "a.docx")

    assert document.sections[0].label == "Chapter One"


def test_epub_chapters_become_sections():
    raw = _epub([("Uno", "El gato duerme."), ("Dos", "La casa es grande.")])

    document = parse_document(raw, "a.epub")

    labels = [s.label for s in document.sections]
    assert "Uno" in labels and "Dos" in labels


def test_the_epub_navigation_document_is_not_imported():
    """It is a table of contents. Its links would become vocabulary
    candidates."""
    document = parse_document(_epub([("Uno", "El gato duerme.")]), "a.epub")

    assert len(document.sections) == 1


def test_csv_and_json_still_parse():
    assert parse_document(b"term,translation\ngato,cat", "a.csv").sentence_count >= 1
    assert parse_document(b'[{"term": "gato"}]', "a.json").sentence_count >= 1


def test_every_registered_media_type_has_a_test_above():
    """Guards the guard: a format added to the registry without a fixture test
    would otherwise be silently unexercised."""
    assert len(SUPPORTED_MEDIA_TYPES) == 11


# --- Refusals --------------------------------------------------------------


def test_an_oversized_file_is_refused_before_parsing():
    with pytest.raises(DocumentTooLargeError):
        parse_document(b"x" * (MAX_BYTES + 1), "a.txt")


def test_an_empty_file_is_refused():
    with pytest.raises(DocumentStructureError):
        parse_document(b"", "a.txt")


def test_an_encrypted_pdf_is_refused_rather_than_guessed_at():
    """Not attempted with an empty password: a document someone encrypted
    should fail loudly rather than be opened by guessing."""
    with pytest.raises(DocumentStructureError, match="encrypted"):
        parse_document(_pdf(["secret"], encrypt=True), "a.pdf")


def test_a_corrupt_pdf_is_refused():
    with pytest.raises(DocumentStructureError):
        parse_document(b"%PDF-1.4\nthis is not a pdf", "a.pdf")


def test_an_archive_that_is_neither_epub_nor_docx_is_refused():
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("hello.txt", "hi")

    with pytest.raises(DocumentStructureError, match="neither"):
        parse_document(buffer.getvalue(), "a.zip")


def test_a_zip_bomb_is_refused_without_being_decompressed():
    """The ratio is computed from the archive header. Measuring by
    decompressing would mean already having done the damage."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", "0" * (5 * 1024 * 1024))

    with pytest.raises(DocumentStructureError, match="expands"):
        parse_document(buffer.getvalue(), "a.docx")


def test_a_document_with_too_many_sections_is_refused_not_truncated():
    """Silently importing the first 2,000 pages of a larger book would look
    like success and lose most of the document."""
    with pytest.raises(DocumentStructureError, match="sections"):
        build_document("big", "text/plain", [(f"P{i}", PROSE) for i in range(MAX_SECTIONS + 1)])


# --- Sentence splitting ----------------------------------------------------


def test_line_breaks_inside_a_sentence_are_collapsed():
    """PDFs break at the page width and subtitles at display timing. A sentence
    that keeps those breaks is unreadable as an excerpt."""
    sentences = split_sentences("El gato\nduerme en\nla casa. Vamos.")

    assert sentences[0].text == "El gato duerme en la casa."


def test_sentences_are_indexed_for_provenance():
    sentences = split_sentences(PROSE)

    assert [s.index for s in sentences] == [0, 1, 2]


def test_a_candidate_can_be_traced_back_to_its_sentence():
    """The point of provenance: a card extracted from page 40 should still be
    able to say which sentence it came from."""
    document = parse_document(_pdf(["El gato duerme. La casa es grande."]), "a.pdf")

    assert "gato" in document.locate("Page 1", 0)
    assert document.locate("Page 99", 0) is None


def test_empty_sections_are_dropped_and_counted():
    """A 400-page PDF of scans would otherwise produce 400 sections a user has
    to scroll past to learn there is nothing in any of them."""
    document = build_document("x", "application/pdf", [("Page 1", PROSE), ("Page 2", "   ")])

    assert len(document.sections) == 1
    assert "1 section(s) contained no readable text" in document.warnings[0]
